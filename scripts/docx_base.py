"""
Infraestrutura de composição do Documento Técnico em .docx.

O conteúdo descreve o sistema tal como ele existe no repositório: cada
afirmação técnica corresponde a código verificado, e não a intenção de
projeto. Manter a geração em script permite regenerar o documento sempre que
a aplicação evoluir.

Reúne o objeto do documento e as funções de composição — parágrafos, títulos,
figuras e tabelas — usadas por gerar-documento.py.
"""

import os

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

FIGURAS = os.path.join('docs', 'academico', 'figuras')
SAIDA = os.path.join('docs', 'academico', 'Documento Tecnico - Fleet Manager.docx')

doc = Document()


# ---------------------------------------------------------------------------
# Estilos gerais — aproximação das normas de trabalho acadêmico
# ---------------------------------------------------------------------------
def configurar():
    s = doc.sections[0]
    s.page_height, s.page_width = Cm(29.7), Cm(21)
    s.top_margin, s.left_margin = Cm(3), Cm(3)
    s.bottom_margin, s.right_margin = Cm(2), Cm(2)

    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for nome, tam in (('Heading 1', 14), ('Heading 2', 12), ('Heading 3', 12)):
        e = doc.styles[nome]
        e.font.name = 'Times New Roman'
        e.font.size = Pt(tam)
        e.font.bold = True
        e.font.color.rgb = RGBColor(0, 0, 0)
        e.paragraph_format.space_before = Pt(18)
        e.paragraph_format.space_after = Pt(12)
        e.paragraph_format.line_spacing = 1.5


def p(texto='', alinhamento=None, negrito=False, tamanho=None, italico=False,
      espaco_antes=None, espaco_depois=None, recuo=True):
    par = doc.add_paragraph()
    r = par.add_run(texto)
    r.bold = negrito
    r.italic = italico
    if tamanho:
        r.font.size = Pt(tamanho)
    if alinhamento is not None:
        par.alignment = alinhamento
    if recuo and par.alignment in (None, WD_ALIGN_PARAGRAPH.JUSTIFY):
        par.paragraph_format.first_line_indent = Cm(1.25)
    if espaco_antes:
        par.paragraph_format.space_before = Pt(espaco_antes)
    if espaco_depois:
        par.paragraph_format.space_after = Pt(espaco_depois)
    return par


def centro(texto='', negrito=False, tamanho=None, espaco_antes=None):
    return p(texto, WD_ALIGN_PARAGRAPH.CENTER, negrito, tamanho,
             espaco_antes=espaco_antes, recuo=False)


def h1(texto):
    doc.add_heading(texto, 1)


def h2(texto):
    doc.add_heading(texto, 2)


def item(texto, nivel=0):
    par = doc.add_paragraph(texto, style='List Bullet' if nivel == 0 else 'List Bullet 2')
    par.paragraph_format.line_spacing = 1.5
    par.paragraph_format.space_after = Pt(0)
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return par


def figura(arquivo, legenda, largura=15.5):
    caminho = os.path.join(FIGURAS, arquivo)
    if not os.path.exists(caminho):
        raise FileNotFoundError(f'Figura ausente: {caminho}. Execute gerar-diagramas.py antes.')
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_before = Pt(12)
    par.add_run().add_picture(caminho, width=Cm(largura))
    leg = doc.add_paragraph()
    leg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    leg.paragraph_format.space_after = Pt(12)
    r = leg.add_run(legenda)
    r.font.size = Pt(10)
    r = leg.add_run('\nFonte: elaborado pelos autores (2026).')
    r.font.size = Pt(10)


def tabela(titulo, cabecalho, linhas, larguras=None):
    leg = doc.add_paragraph()
    leg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    leg.paragraph_format.space_before = Pt(12)
    r = leg.add_run(titulo)
    r.font.size = Pt(10)

    t = doc.add_table(rows=1, cols=len(cabecalho))
    t.style = 'Table Grid'
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, c in enumerate(cabecalho):
        cel = t.rows[0].cells[i]
        cel.text = ''
        par = cel.paragraphs[0]
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.paragraph_format.line_spacing = 1.0
        run = par.add_run(c)
        run.bold = True
        run.font.size = Pt(10)
    for ln in linhas:
        cels = t.add_row().cells
        for i, v in enumerate(ln):
            cels[i].text = ''
            par = cels[i].paragraphs[0]
            par.paragraph_format.line_spacing = 1.0
            par.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = par.add_run(str(v))
            run.font.size = Pt(10)
    if larguras:
        for row in t.rows:
            for i, w in enumerate(larguras):
                row.cells[i].width = Cm(w)

    fonte = doc.add_paragraph()
    fonte.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fonte.paragraph_format.space_after = Pt(12)
    r = fonte.add_run('Fonte: elaborado pelos autores (2026).')
    r.font.size = Pt(10)
    return t


def quebra():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def sumario_automatico():
    """Insere um campo de sumário que o Word preenche ao atualizar."""
    par = doc.add_paragraph()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')
    filho = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = ('Para exibir o sumário: clique com o botão direito aqui e escolha '
              '"Atualizar campo".')
    filho.append(t)
    fld.append(filho)
    par._p.append(fld)
